import fitz
from docx import Document
import os


def extract_text(file_path: str):

    extension = os.path.splitext(file_path)[1].lower()

    if extension == ".pdf":
        return extract_pdf(file_path)

    elif extension == ".docx":
        return extract_docx(file_path)

    elif extension == ".txt":
        return extract_txt(file_path)

    else:
        raise Exception("Unsupported file type")


def extract_pdf(file_path):

    text = ""

    pdf = fitz.open(file_path)

    for page in pdf:
        text += page.get_text()

    pdf.close()

    return text


def extract_docx(file_path):

    text = ""

    document = Document(file_path)

    for paragraph in document.paragraphs:
        text += paragraph.text + "\n"

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                text += cell.text + " "
            text += "\n"

    return text


def extract_txt(file_path):
    with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()